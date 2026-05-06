r"""
book_build.py — Tum bolumleri birlestirip DOCX/PDF ciktisi uretir.

Kullanim:
  .venv\Scripts\python.exe tools/book_build.py
  .venv\Scripts\python.exe tools/book_build.py --format docx
  .venv\Scripts\python.exe tools/book_build.py --format md
"""

import os
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent.parent
BOOK_PROJECT = ROOT / "book_projects" / "java-temelleri"
sys.path.insert(0, str(ROOT))

# Bölüm sırası
CHAPTER_ORDER = [
    "bolum-01", "bolum-02", "bolum-03", "bolum-04", "bolum-05", "bolum-06",
    "bolum-07", "bolum-08", "bolum-09", "bolum-10", "bolum-11",
    "bolum-12", "bolum-13", "bolum-14", "bolum-15", "bolum-16",
    "bolum-17", "bolum-18", "bolum-19", "bolum-20", "bolum-21",
    "bolum-22", "bolum-23",
    "ek-a", "ek-b", "ek-c", "ek-d",
]

# Bölüm başlıkları (progress.json'dan alınan)
CHAPTER_TITLES = {
    "bolum-01": "Java'ya Giriş ve Temel Kavramlar",
    "bolum-02": "Değişkenler, Veri Tipleri ve Tip Dönüşümleri",
    "bolum-03": "Operatörler ve İfadeler",
    "bolum-04": "Kontrol Akışı: Karar Yapıları",
    "bolum-05": "Döngüler ve Tekrarlı İşlemler",
    "bolum-06": "Metotlar ve Fonksiyonel Programlama",
    "bolum-07": "Algoritmik Problem Çözme Desenleri",
    "bolum-08": "Metotlar, Overloading ve Özyineleme",
    "bolum-09": "Diziler ve Çok Boyutlu Veri Yapıları",
    "bolum-10": "String İşlemleri ve Metin Problemleri",
    "bolum-11": "Matematiksel Yardımcılar ve Rastgelelik",
    "bolum-12": "Tarih ve Zaman İşlemleri",
    "bolum-13": "Paketler, import Kullanımı ve Proje Düzeni",
    "bolum-14": "Koleksiyonlar ve Dinamik Veri Yönetimi",
    "bolum-15": "Hata Yönetimi ve Dayanıklı Programlama",
    "bolum-16": "Dosya İşlemleri ve Kalıcı Veri Saklama",
    "bolum-17": "Sınıf, Nesne, Constructor ve Kapsülleme",
    "bolum-18": "Kalıtım ve Interface'e Kısa Bir Bakış",
    "bolum-19": "GUI Programlamaya Giriş ve Swing Arayüz Tasarımı",
    "bolum-20": "Temel Swing Bileşenleri, Olay Yönetimi ve Form Doğrulama",
    "bolum-21": "Liste, Tablo, Menu ve Diyaloglarla GUI Veri Sunumu",
    "bolum-22": "JDBC ile Veritabanı Programlamaya Giriş",
    "bolum-23": "Bütünleşik Uygulama ve Final Proje Rehberi",
    "ek-a": "Ek A: Sık Yapılan Java Hataları ve Çözüm Rehberi",
    "ek-b": "Ek B: JavaFX'e Kısa Bakış",
    "ek-c": "Ek C: Mini Proje Fikirleri ve Rubrikler",
    "ek-d": "Ek D: Java Programlama Kontrol Rehberi, Sık Hatalar ve Kod Kalitesi",
}


def strip_frontmatter(text):
    """YAML front matter'ı kaldırır."""
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            return text[end + 3:].strip()
    return text


def extract_title_from_frontmatter(text):
    """Front matter'dan başlık alır."""
    m = re.match(r"---\s*\ntitle:\s*\"(.+?)\"", text)
    if m:
        return m.group(1)
    m = re.match(r"---\s*\ntitle:\s*(.+?)\s*\n", text)
    if m:
        return m.group(1).strip()
    return None


def read_chapter(chapter_id):
    """Bölüm dosyasını okur, front matter'ı çıkarır, bölüm başlığı ekler."""
    path = BOOK_PROJECT / "chapters" / chapter_id / "draft_versions" / "v001.md"
    if not path.exists():
        print(f"  [UYARI] {chapter_id} bulunamadi: {path}")
        return None

    content = path.read_text(encoding="utf-8")
    title = CHAPTER_TITLES.get(chapter_id, chapter_id)

    # Front matter'ı kaldır
    body = strip_frontmatter(content)

    # İlk H1'i kontrol et — yoksa ekle
    if not re.search(r"^#\s", body, re.MULTILINE):
        body = f"# {title}\n\n{body}"

    return body


def merge_book():
    """Tüm bölümleri birleştirir."""
    print("Kitap birleştiriliyor...")
    print(f"  Bölüm sayısı: {len(CHAPTER_ORDER)}")

    chapters = []
    total_chars = 0
    for i, ch_id in enumerate(CHAPTER_ORDER, 1):
        body = read_chapter(ch_id)
        if body:
            chapters.append(body)
            total_chars += len(body)
            print(f"  [{i:02d}/{len(CHAPTER_ORDER)}] {ch_id} — {len(body):,} chars")
        else:
            print(f"  [{i:02d}/{len(CHAPTER_ORDER)}] {ch_id} — ATLANDI")

    # Başlık sayfası
    book_title = """---
title: "Java Programlamaya Giriş"
subtitle: "Temel Kavramlardan İleri Uygulamalara"
author: "bookMaker AI"
date: "2026-05-04"
lang: "tr"
---

# Java Programlamaya Giriş

**Temel Kavramlardan İleri Uygulamalara**

---

*Bu kitap, Java programlama dilini sıfırdan öğrenmek isteyenler için hazırlanmıştır.
Temel kavramlardan başlayarak, nesne yönelimli programlama, GUI geliştirme,
veritabanı işlemleri ve ileri düzey konulara kadar kapsamlı bir içerik sunar.*

---

"""

    full_text = book_title + "\n\n".join(chapters)
    print(f"\n  Toplam: {total_chars:,} chars")
    print(f"  Birleşik: {len(full_text):,} chars")
    return full_text


def save_md(text, path):
    """Markdown olarak kaydet."""
    path.write_text(text, encoding="utf-8")
    print(f"\n[OK] Markdown kaydedildi: {path} ({len(text):,} chars)")


def save_docx(text, path):
    """python-docx ile DOCX oluştur."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("\n❌ python-docx kurulu değil. Önce yükleyin: pip install python-docx")
        return

    doc = Document()

    # Varsayılan stil
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Satırları işle
    lines = text.split('\n')
    in_code_block = False
    code_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Kod bloğu
        if line.startswith('```'):
            if in_code_block:
                # Kod bloğu bitişi
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                # Dil etiketini atla
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # CODE_META satırlarını atla
        if line.strip().startswith('<!-- CODE_META') or line.strip().startswith('-->'):
            i += 1
            continue

        # Başlıklar
        if line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.left_indent = Inches(0.3)
        elif line.strip() == '---':
            doc.add_paragraph('_' * 60)
        elif line.strip() == '':
            doc.add_paragraph('')
        else:
            # Normal paragraf
            p = doc.add_paragraph(line)

        i += 1

    doc.save(str(path))
    print(f"\n[OK] DOCX kaydedildi: {path}")


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Kitap birleştirme ve dönüşüm")
    parser.add_argument("--format", choices=["md", "docx", "both"], default="both",
                        help="Çıktı formatı (varsayılan: both)")
    args = parser.parse_args()

    os.makedirs(BOOK_PROJECT / "build" / "output", exist_ok=True)

    # Birleştir
    text = merge_book()

    # MD çıktı
    md_path = BOOK_PROJECT / "build" / "output" / "java-programlamaya-giris.md"
    save_md(text, md_path)

    # DOCX çıktı
    if args.format in ("docx", "both"):
        docx_path = BOOK_PROJECT / "build" / "output" / "java-programlamaya-giris.docx"
        save_docx(text, docx_path)


if __name__ == "__main__":
    main()
